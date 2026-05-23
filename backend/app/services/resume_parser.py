from io import BytesIO

from docx import Document
from pypdf import PdfReader


PDF_CONTENT_TYPES = {"application/pdf"}
DOCX_CONTENT_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
TXT_CONTENT_TYPES = {"text/plain"}


class ResumeParseError(Exception):
    """Raised when a resume file cannot be parsed into text."""


class UnsupportedResumeContentTypeError(ResumeParseError):
    """Raised when the uploaded resume content type is not supported."""


def _normalize_content_type(content_type: str) -> str:
    return content_type.split(";", maxsplit=1)[0].strip().lower()


def _normalize_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    meaningful_lines = [line for line in lines if line]

    return "\n".join(meaningful_lines).strip()


def _ensure_text_found(text: str, file_type: str) -> str:
    normalized_text = _normalize_text(text)

    if not normalized_text:
        raise ResumeParseError(
            f"No extractable text was found in the {file_type} resume.",
        )

    return normalized_text


def parse_resume_file(file_bytes: bytes, content_type: str) -> str:
    normalized_content_type = _normalize_content_type(content_type)

    if normalized_content_type in PDF_CONTENT_TYPES:
        return parse_pdf(file_bytes)

    if normalized_content_type in DOCX_CONTENT_TYPES:
        return parse_docx(file_bytes)

    if normalized_content_type in TXT_CONTENT_TYPES:
        return parse_txt(file_bytes)

    raise UnsupportedResumeContentTypeError(
        f"Unsupported resume content type: {content_type}.",
    )


def parse_pdf(file_bytes: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(file_bytes))
        page_text = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:
        raise ResumeParseError("Unable to parse PDF resume.") from exc

    return _ensure_text_found("\n".join(page_text), "PDF")


def parse_docx(file_bytes: bytes) -> str:
    try:
        document = Document(BytesIO(file_bytes))
        paragraph_text = [paragraph.text for paragraph in document.paragraphs]
        table_text = [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
    except Exception as exc:
        raise ResumeParseError("Unable to parse DOCX resume.") from exc

    return _ensure_text_found("\n".join([*paragraph_text, *table_text]), "DOCX")


def parse_txt(file_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            return _ensure_text_found(file_bytes.decode(encoding), "TXT")
        except UnicodeDecodeError:
            continue

    raise ResumeParseError("Unable to parse TXT resume.")
